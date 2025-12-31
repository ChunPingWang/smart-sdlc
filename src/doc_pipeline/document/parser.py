"""Multi-format document parser."""

import re
import warnings
from pathlib import Path
from typing import Any

from doc_pipeline.document.models import (
    DocumentChunk,
    DocumentMetadata,
    ParsedDocument,
)


# DDL 相關關鍵字
DDL_KEYWORDS = {
    "資料庫", "資料表", "數據庫", "數據表", "database", "table", "schema",
    "欄位", "字段", "column", "field", "primary key", "foreign key",
    "pk", "fk", "index", "constraint", "資料型別", "資料類型", "data type",
}


class DocumentParser:
    """Parser for multiple document formats (Word, Excel, Markdown, PDF)."""

    SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".md", ".markdown", ".pdf", ".txt"}

    def __init__(self) -> None:
        """Initialize the document parser."""
        pass

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        Parse a document file and return structured content.

        Args:
            file_path: Path to the document file

        Returns:
            ParsedDocument containing parsed content and metadata
        """
        file_path = Path(file_path)
        self._validate_file(file_path)

        ext = file_path.suffix.lower()

        if ext in {".md", ".markdown", ".txt"}:
            elements = self._parse_markdown(file_path)
        elif ext == ".docx":
            elements = self._parse_docx(file_path)
        elif ext == ".xlsx":
            elements = self._parse_xlsx(file_path)
        elif ext == ".pdf":
            elements = self._parse_pdf(file_path)
        else:
            elements = []

        chunks = self._create_chunks(elements, file_path)

        title = ""
        for elem in elements:
            if elem.get("type") == "Title":
                title = elem.get("text", "")
                break

        return ParsedDocument(
            source_file=file_path,
            file_type=file_path.suffix.lstrip("."),
            title=title,
            chunks=chunks,
            raw_elements=elements,
        )

    def _parse_markdown(self, file_path: Path) -> list[dict[str, Any]]:
        """Parse markdown file."""
        content = file_path.read_text(encoding="utf-8")
        elements = []

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                elements.append({"type": "Title", "text": line[2:]})
            elif line.startswith("## "):
                elements.append({"type": "Header", "text": line[3:]})
            elif line.startswith("### "):
                elements.append({"type": "Header", "text": line[4:]})
            elif line.startswith("- ") or line.startswith("* "):
                elements.append({"type": "ListItem", "text": line[2:]})
            elif re.match(r"^\d+\.", line):
                elements.append({"type": "ListItem", "text": re.sub(r"^\d+\.\s*", "", line)})
            elif line.startswith("|"):
                elements.append({"type": "Table", "text": line})
            elif line.startswith("```"):
                elements.append({"type": "CodeBlock", "text": line})
            else:
                elements.append({"type": "NarrativeText", "text": line})

        return elements

    def _parse_docx(self, file_path: Path) -> list[dict[str, Any]]:
        """Parse Word document."""
        try:
            from docx import Document
        except ImportError:
            return [{"type": "Error", "text": "python-docx not installed"}]

        doc = Document(str(file_path))
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""

            if "Heading 1" in style or "Title" in style:
                elements.append({"type": "Title", "text": text})
            elif "Heading" in style:
                elements.append({"type": "Header", "text": text})
            else:
                elements.append({"type": "NarrativeText", "text": text})

        return elements

    def _parse_xlsx(self, file_path: Path) -> list[dict[str, Any]]:
        """Parse Excel file."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return [{"type": "Error", "text": "openpyxl not installed"}]

        wb = load_workbook(file_path, read_only=True)
        elements = []

        for sheet in wb.sheetnames:
            ws = wb[sheet]
            elements.append({"type": "Title", "text": f"Sheet: {sheet}"})

            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                if row_text.strip(" |"):
                    elements.append({"type": "Table", "text": row_text})

        return elements

    def _parse_pdf(self, file_path: Path) -> list[dict[str, Any]]:
        """Parse PDF file."""
        try:
            import pdfplumber
        except ImportError:
            return [{"type": "Error", "text": "pdfplumber not installed"}]

        elements = []

        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    elements.append({
                        "type": "NarrativeText",
                        "text": text,
                        "metadata": {"page_number": i + 1},
                    })

        return elements

    def parse_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
    ) -> list[ParsedDocument]:
        """
        Parse all supported documents in a directory.

        Args:
            directory: Path to the directory
            recursive: Whether to search recursively

        Returns:
            List of ParsedDocument objects
        """
        directory = Path(directory)
        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.parse(file_path)
                    documents.append(doc)
                except Exception as e:
                    print(f"Error parsing {file_path}: {e}")

        return documents

    def _validate_file(self, file_path: Path) -> None:
        """Validate the file exists and is supported."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

    def _create_chunks(
        self,
        elements: list[dict[str, Any]],
        source_file: Path,
    ) -> list[DocumentChunk]:
        """Create document chunks from parsed elements."""
        chunks = []
        current_section = ""
        current_chapter = ""

        for elem in elements:
            elem_type = elem.get("type", "")
            text = elem.get("text", "").strip()

            if not text:
                continue

            if elem_type == "Title":
                current_chapter = text
                current_section = ""
            elif elem_type == "Header":
                current_section = text

            metadata = DocumentMetadata(
                source_file=source_file,
                chapter=current_chapter,
                section=current_section,
            )

            if "metadata" in elem and "page_number" in elem["metadata"]:
                metadata.page_or_section = f"page_{elem['metadata']['page_number']}"
                metadata.extra["page_number"] = elem["metadata"]["page_number"]

            chunk = DocumentChunk(
                content=text,
                metadata=metadata,
            )
            chunks.append(chunk)

        return chunks


def parse_file(file_path: str | Path) -> ParsedDocument:
    """Convenience function to parse a single file."""
    parser = DocumentParser()
    return parser.parse(file_path)


def parse_directory(
    directory: str | Path,
    recursive: bool = True,
) -> list[ParsedDocument]:
    """Convenience function to parse a directory."""
    parser = DocumentParser()
    return parser.parse_directory(directory, recursive)


# =============================================================================
# Conversion Functions
# =============================================================================


def convert_to_markdown(file_path: str | Path, use_docling: bool = False) -> str:
    """
    Convert a document file directly to Markdown format without any modification.

    Args:
        file_path: Path to the document file
        use_docling: Use Docling for AI-powered conversion (more accurate)

    Returns:
        Markdown string representation of the document
    """
    parser = DocumentParser()
    file_path = Path(file_path)
    parser._validate_file(file_path)

    ext = file_path.suffix.lower()

    if use_docling:
        return _convert_with_docling(file_path)

    if ext in {".md", ".markdown", ".txt"}:
        return file_path.read_text(encoding="utf-8")
    elif ext == ".docx":
        return _docx_to_markdown(file_path)
    elif ext == ".xlsx":
        # For single file conversion, combine all sheets
        sheets = _xlsx_to_markdown_sheets(file_path)
        return "\n\n---\n\n".join(sheets.values())
    elif ext == ".pdf":
        return _pdf_to_markdown(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _convert_with_docling(file_path: Path) -> str:
    """Convert document using Docling AI-powered converter."""
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        raise ImportError(
            "Docling not installed. Run: pip install docling\n"
            "Docling provides AI-powered document conversion with better accuracy."
        )

    converter = DocumentConverter()
    result = converter.convert(str(file_path))
    return result.document.export_to_markdown()


def _docx_to_markdown(file_path: Path) -> str:
    """Convert Word document to Markdown with proper heading hierarchy."""
    try:
        from docx import Document
        from docx.table import Table as DocxTable
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(file_path))
    lines: list[str] = []
    list_counter = 0
    in_list = False

    def get_heading_level(style_name: str) -> int | None:
        """Extract heading level from style name."""
        if not style_name:
            return None
        if "Title" in style_name:
            return 1
        match = re.search(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            return int(match.group(1))
        if "Heading" in style_name:
            return 2
        return None

    def format_table(table: DocxTable) -> list[str]:
        """Convert Word table to Markdown table."""
        table_lines = []
        rows_data = []

        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_data.append(row_cells)

        if not rows_data:
            return []

        # Determine column count
        max_cols = max(len(row) for row in rows_data)

        # Header row
        header = rows_data[0]
        header.extend([""] * (max_cols - len(header)))
        table_lines.append("| " + " | ".join(header) + " |")
        table_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        # Data rows
        for row in rows_data[1:]:
            row.extend([""] * (max_cols - len(row)))
            table_lines.append("| " + " | ".join(row) + " |")

        return table_lines

    # Process document body (paragraphs and tables in order)
    for element in doc.element.body:
        # Check if it's a paragraph
        if element.tag.endswith("p"):
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    style_name = para.style.name if para.style else ""

                    # Handle empty paragraphs
                    if not text:
                        if lines and lines[-1] != "":
                            lines.append("")
                        in_list = False
                        list_counter = 0
                        continue

                    # Check heading level
                    heading_level = get_heading_level(style_name)
                    if heading_level:
                        if lines and lines[-1] != "":
                            lines.append("")
                        lines.append("#" * heading_level + " " + text)
                        lines.append("")
                        in_list = False
                        list_counter = 0
                        continue

                    # Check for list items
                    if "List" in style_name:
                        if "Number" in style_name or "Ordered" in style_name:
                            list_counter += 1
                            lines.append(f"{list_counter}. {text}")
                        else:
                            lines.append(f"- {text}")
                        in_list = True
                        continue

                    # Check for indented text (TOC or nested content)
                    if "TOC" in style_name:
                        # Table of contents - add as indented list
                        indent_level = 0
                        if para.paragraph_format.left_indent:
                            indent_level = int(para.paragraph_format.left_indent.pt / 36)
                        lines.append("  " * indent_level + "- " + text)
                        continue

                    # Regular paragraph
                    if in_list:
                        lines.append("")
                        in_list = False
                        list_counter = 0

                    lines.append(text)
                    lines.append("")
                    break

        # Check if it's a table
        elif element.tag.endswith("tbl"):
            for table in doc.tables:
                if table._element is element:
                    if lines and lines[-1] != "":
                        lines.append("")
                    table_lines = format_table(table)
                    lines.extend(table_lines)
                    lines.append("")
                    break

    return "\n".join(lines)


def _xlsx_to_markdown_sheets(file_path: Path) -> dict[str, str]:
    """
    Convert Excel file to Markdown, one file per sheet.

    Returns:
        Dictionary mapping sheet names to their Markdown content
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    sheets_content: dict[str, str] = {}

    with warnings.catch_warnings():
        warnings.filterwarnings("ignore", category=UserWarning)
        wb = load_workbook(file_path, read_only=True, data_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines = [f"# {sheet_name}", ""]

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                sheets_content[sheet_name] = "\n".join(lines)
                continue

            # Filter out completely empty rows
            non_empty_rows = [row for row in rows if any(cell is not None for cell in row)]
            if not non_empty_rows:
                sheets_content[sheet_name] = "\n".join(lines)
                continue

            # Determine max columns
            max_cols = max(len(row) for row in non_empty_rows)

            # Create table header from first row
            header_row = non_empty_rows[0]
            header_cells = [str(cell) if cell is not None else "" for cell in header_row]
            header_cells.extend([""] * (max_cols - len(header_cells)))

            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")

            # Data rows
            for row in non_empty_rows[1:]:
                cells = [str(cell) if cell is not None else "" for cell in row]
                cells.extend([""] * (max_cols - len(cells)))
                lines.append("| " + " | ".join(cells) + " |")

            lines.append("")
            sheets_content[sheet_name] = "\n".join(lines)

    return sheets_content


def _xlsx_to_markdown(file_path: Path) -> str:
    """Convert Excel file to single Markdown (legacy, combines all sheets)."""
    sheets = _xlsx_to_markdown_sheets(file_path)
    return "\n\n---\n\n".join(sheets.values())


def _pdf_to_markdown(file_path: Path) -> str:
    """Convert PDF to Markdown."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

    lines = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            if i > 0:
                lines.append("")
                lines.append("---")
                lines.append("")

            lines.append(f"## Page {i + 1}")
            lines.append("")

            text = page.extract_text()
            if text:
                lines.append(text)

    return "\n".join(lines)


# =============================================================================
# DDL Generation Functions
# =============================================================================


def contains_ddl_keywords(content: str) -> bool:
    """Check if content contains DDL-related keywords."""
    content_lower = content.lower()
    return any(keyword.lower() in content_lower for keyword in DDL_KEYWORDS)


def generate_ddl_from_markdown(markdown_content: str, table_name: str | None = None) -> str:
    """
    Generate DDL SQL from Markdown content.

    Supports FSD (Functional Specification Document) table format:
    - Row with "資料表名稱(英)" contains the table name
    - Row with "欄位名稱(中)" and "欄位名稱(英)" is the header
    - Following rows are column definitions

    Args:
        markdown_content: Markdown content containing table definitions
        table_name: Optional fallback table name

    Returns:
        DDL SQL statements
    """
    lines = markdown_content.split("\n")

    # Try to extract FSD format table
    result = _extract_fsd_table(lines, table_name)

    if result:
        return result

    return ""


def _extract_fsd_table(lines: list[str], fallback_name: str | None = None) -> str:
    """
    Extract table definition from FSD (Functional Specification Document) format.

    FSD Format:
    |  | 資料表名稱(中) | 個簽優惠資料檔 | ... |
    | --- | --- | --- | ... |
    |  | 資料表名稱(英) | CUST_PREFERENTIAL_MST | ... |
    |  | 欄位名稱(中) | 欄位名稱(英) | 欄位型別 | 長度 | PK | UK | 必填 | 預設值 | 說明 |
    |  | 流水號 | ID | INTEGER | 30 | Y |  |  |  | 系統生成的流水號 |
    """
    table_name = None
    columns: list[dict[str, Any]] = []
    in_column_section = False
    header_indices: dict[str, int] = {}

    for line in lines:
        line = line.strip()
        if not line.startswith("|") or "---" in line:
            continue

        cells = [c.strip() for c in line.split("|")]
        # Remove empty first and last cells from split
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]

        if len(cells) < 3:
            continue

        # Check for table name row: "資料表名稱(英)"
        if "資料表名稱(英)" in cells[1] or "資料表名稱（英）" in cells[1]:
            # Table name is in the next cell
            potential_name = cells[2].strip()
            if potential_name and re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', potential_name):
                table_name = potential_name
            continue

        # Check for column header row
        if ("欄位名稱(中)" in cells[1] or "欄位名稱（中）" in cells[1]) and \
           ("欄位名稱(英)" in cells[2] or "欄位名稱（英）" in cells[2]):
            in_column_section = True
            # Map header positions
            for i, cell in enumerate(cells):
                cell_lower = cell.lower().replace("（", "(").replace("）", ")")
                if "欄位名稱(英)" in cell or "column" in cell_lower:
                    header_indices["column_name"] = i
                elif "欄位型別" in cell or "type" in cell_lower:
                    header_indices["column_type"] = i
                elif "長度" in cell or "length" in cell_lower:
                    header_indices["length"] = i
                elif cell.upper() == "PK" or "primary" in cell_lower:
                    header_indices["pk"] = i
                elif cell.upper() == "UK" or "unique" in cell_lower:
                    header_indices["uk"] = i
                elif "必填" in cell or "required" in cell_lower or "not null" in cell_lower:
                    header_indices["required"] = i
                elif "說明" in cell or "description" in cell_lower or "comment" in cell_lower:
                    header_indices["comment"] = i
            continue

        # Parse column definition row
        if in_column_section and len(cells) >= 4:
            col_info = _parse_fsd_column_row(cells, header_indices)
            if col_info:
                columns.append(col_info)

    # Generate DDL if we have valid data
    if not columns:
        return ""

    # Use fallback name if no table name found
    if not table_name:
        if fallback_name:
            # Extract English table name from fallback if possible
            match = re.search(r'([A-Z][A-Z0-9_]+)', fallback_name)
            if match:
                table_name = match.group(1)
            else:
                table_name = re.sub(r'[^\w]', '_', fallback_name)
        else:
            return ""

    # Build DDL
    ddl_lines = [
        f"-- Auto-generated DDL from FSD document",
        f"-- Table: {table_name}",
        "",
        f"CREATE TABLE {table_name} ("
    ]

    col_defs = []
    for col in columns:
        col_def = f"    {col['name']} {col['type']}"
        if col.get("constraints"):
            col_def += " " + " ".join(col["constraints"])
        if col.get("comment"):
            # Escape comment for SQL
            comment = col["comment"].replace("'", "''").replace("\n", " ")
            col_def += f"  -- {comment}"
        col_defs.append(col_def)

    ddl_lines.append(",\n".join(col_defs))
    ddl_lines.append(");")
    ddl_lines.append("")

    return "\n".join(ddl_lines)


def _parse_fsd_column_row(cells: list[str], header_indices: dict[str, int]) -> dict[str, Any] | None:
    """Parse a column definition row from FSD format."""
    # Get column name (English)
    col_name_idx = header_indices.get("column_name", 2)
    if col_name_idx >= len(cells):
        return None

    col_name = cells[col_name_idx].strip()

    # Skip if empty or looks like a header
    if not col_name or col_name.lower() in ["欄位名稱(英)", "column", "field", ""]:
        return None

    # Validate column name (must be valid SQL identifier)
    if not re.match(r'^[A-Za-z_][A-Za-z0-9_]*$', col_name):
        return None

    # Get column type
    col_type_idx = header_indices.get("column_type", 3)
    raw_type = cells[col_type_idx].strip().upper() if col_type_idx < len(cells) else "VARCHAR"

    # Get length
    length_idx = header_indices.get("length", 4)
    length = cells[length_idx].strip() if length_idx < len(cells) else ""

    # Build SQL type
    col_type = _build_sql_type(raw_type, length)

    # Get constraints
    constraints = []

    # Check PK
    pk_idx = header_indices.get("pk", 5)
    if pk_idx < len(cells):
        pk_value = cells[pk_idx].strip().upper()
        if pk_value in ["Y", "V", "YES", "TRUE", "1", "○", "●"]:
            constraints.append("PRIMARY KEY")

    # Check UK (Unique Key)
    uk_idx = header_indices.get("uk", 6)
    if uk_idx < len(cells):
        uk_value = cells[uk_idx].strip().upper()
        if uk_value in ["Y", "V", "YES", "TRUE", "1", "○", "●"]:
            constraints.append("UNIQUE")

    # Check Required (NOT NULL)
    req_idx = header_indices.get("required", 7)
    if req_idx < len(cells):
        req_value = cells[req_idx].strip().upper()
        if req_value in ["Y", "V", "YES", "TRUE", "1", "○", "●"]:
            constraints.append("NOT NULL")

    # Get comment
    comment_idx = header_indices.get("comment", -1)
    comment = ""
    if comment_idx > 0 and comment_idx < len(cells):
        comment = cells[comment_idx].strip()
    elif len(cells) > 8:
        # Last cell is often the comment
        comment = cells[-1].strip()

    return {
        "name": col_name,
        "type": col_type,
        "constraints": constraints,
        "comment": comment,
    }


def _build_sql_type(raw_type: str, length: str) -> str:
    """Build SQL type with length specification."""
    raw_type = raw_type.upper().strip()
    length = length.strip()

    # Type mapping
    type_map = {
        "VARCHAR": "VARCHAR",
        "VARCHAR2": "VARCHAR",
        "CHAR": "CHAR",
        "NVARCHAR": "NVARCHAR",
        "NVARCHAR2": "NVARCHAR",
        "TEXT": "TEXT",
        "CLOB": "TEXT",
        "INTEGER": "INTEGER",
        "INT": "INTEGER",
        "BIGINT": "BIGINT",
        "SMALLINT": "SMALLINT",
        "NUMBER": "DECIMAL",
        "NUMERIC": "DECIMAL",
        "DECIMAL": "DECIMAL",
        "FLOAT": "FLOAT",
        "DOUBLE": "DOUBLE",
        "REAL": "REAL",
        "DATE": "DATE",
        "DATETIME": "DATETIME",
        "TIMESTAMP": "TIMESTAMP",
        "TIME": "TIME",
        "BOOLEAN": "BOOLEAN",
        "BOOL": "BOOLEAN",
        "BLOB": "BLOB",
        "BINARY": "BINARY",
    }

    sql_type = type_map.get(raw_type, "VARCHAR")

    # Add length for types that need it
    if sql_type in ["VARCHAR", "NVARCHAR", "CHAR"] and length:
        try:
            length_int = int(length)
            return f"{sql_type}({length_int})"
        except ValueError:
            return f"{sql_type}(255)"
    elif sql_type == "DECIMAL" and length:
        # Handle precision,scale format
        if "," in length:
            return f"DECIMAL({length})"
        else:
            try:
                precision = int(length)
                return f"DECIMAL({precision},2)"
            except ValueError:
                return "DECIMAL(18,2)"

    return sql_type
